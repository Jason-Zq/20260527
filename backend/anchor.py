"""
Anchor 抽象 —— Word 模板填值的"位置坐标"模型。

设计目标：用一个不可变的 docx XML 节点引用替代字符串锚点（original_text）。
所有 apply_value 操作都保证：除目标节点（cell/paragraph/run）外，docx 其余
XML 字节级保持不变。

三种 anchor 粒度（精度从粗到细）：
  - AnchorCell      → 整个表格单元格
  - AnchorParagraph → 容器内的某个段落
  - AnchorRun       → 段落内的某个 run，可选 (char_offset, char_length) 字符级

容器类型：
  - ContainerBody   → 文档正文
  - ContainerCell   → 某表格单元格内
  - ContainerHeader → 某 section 的页眉
  - ContainerFooter → 某 section 的页脚

依赖：python-docx 1.x + lxml 4.x+ + pydantic 2.x
"""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal
from typing import Annotated, Any, Literal, Optional, Union

from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from pydantic import BaseModel, Field, model_validator


# ====================== 异常 ======================


class AnchorNotFound(Exception):
    """anchor 在 doc 中找不到对应节点。"""


# ====================== Container 模型 ======================


class ContainerBody(BaseModel):
    kind: Literal["body"]


class ContainerCell(BaseModel):
    kind: Literal["cell"]
    t: int = Field(ge=0, description="table_index")
    r: int = Field(ge=0, description="row_index")
    c: int = Field(ge=0, description="col_index")


class ContainerHeader(BaseModel):
    kind: Literal["header"]
    s: int = Field(ge=0, description="section_index")


class ContainerFooter(BaseModel):
    kind: Literal["footer"]
    s: int = Field(ge=0, description="section_index")


Container = Annotated[
    Union[ContainerBody, ContainerCell, ContainerHeader, ContainerFooter],
    Field(discriminator="kind"),
]


# ====================== Anchor 模型 ======================


class AnchorCell(BaseModel):
    kind: Literal["cell"]
    t: int = Field(ge=0)
    r: int = Field(ge=0)
    c: int = Field(ge=0)


class AnchorParagraph(BaseModel):
    kind: Literal["paragraph"]
    container: Container
    p: int = Field(ge=0, description="paragraph_index in container")


class AnchorRun(BaseModel):
    kind: Literal["run"]
    container: Container
    p: int = Field(ge=0)
    run_index: int = Field(ge=0)
    char_offset: Optional[int] = Field(default=None, ge=0)
    char_length: Optional[int] = Field(default=None, ge=0)

    @model_validator(mode="after")
    def _offset_and_length_together(self) -> "AnchorRun":
        if (self.char_offset is None) != (self.char_length is None):
            raise ValueError(
                "char_offset 和 char_length 必须同时提供或同时省略"
            )
        return self


Anchor = Annotated[
    Union[AnchorCell, AnchorParagraph, AnchorRun],
    Field(discriminator="kind"),
]


# ====================== 定位结果 ======================


@dataclass
class LocatedNode:
    """locate_anchor 的返回值。按 anchor.kind 只有一个字段非空。"""
    cell: Optional[_Cell] = None
    paragraph: Optional[Paragraph] = None
    run: Optional[Run] = None
    slice_range: Optional[tuple[int, int]] = None  # AnchorRun 的 (offset, length)


# ====================== locate_anchor ======================


def _get_container_paragraphs(doc: DocxDocument, container: BaseModel) -> list[Paragraph]:
    """根据 container 拿到段落列表。任何 IndexError → AnchorNotFound。"""
    try:
        if isinstance(container, ContainerBody):
            return list(doc.paragraphs)
        if isinstance(container, ContainerCell):
            cell = doc.tables[container.t].rows[container.r].cells[container.c]
            return list(cell.paragraphs)
        if isinstance(container, ContainerHeader):
            return list(doc.sections[container.s].header.paragraphs)
        if isinstance(container, ContainerFooter):
            return list(doc.sections[container.s].footer.paragraphs)
    except (IndexError, AttributeError) as e:
        raise AnchorNotFound(f"container 定位失败: {container} ({e})") from e
    raise AnchorNotFound(f"未知 container 类型: {type(container)}")


def _is_merged_continuation(cell: _Cell, t_idx: int, r_idx: int, c_idx: int, doc) -> bool:
    """判断 cell 是否是合并 cell 的"被合并"部分（不是左上起点）。

    python-docx 在 row.cells 中对合并 cell 会重复返回同一 _tc 实例，
    所以检测：左侧是否有同 _tc 实例的 cell。
    """
    try:
        row = doc.tables[t_idx].rows[r_idx]
        for prev_c in range(c_idx):
            if id(row.cells[prev_c]._tc) == id(cell._tc):
                return True
    except (IndexError, AttributeError):
        return False
    return False


def locate_anchor(doc: DocxDocument, anchor: BaseModel) -> LocatedNode:
    """根据 anchor 在 docx 内找到目标节点。失败抛 AnchorNotFound。"""
    if isinstance(anchor, AnchorCell):
        try:
            cell = doc.tables[anchor.t].rows[anchor.r].cells[anchor.c]
        except (IndexError, AttributeError) as e:
            raise AnchorNotFound(
                f"cell anchor 越界: t={anchor.t},r={anchor.r},c={anchor.c} ({e})"
            ) from e
        if _is_merged_continuation(cell, anchor.t, anchor.r, anchor.c, doc):
            raise AnchorNotFound(
                f"cell anchor 指向了合并 cell 的被合并位（应指向左上起点）: "
                f"t={anchor.t},r={anchor.r},c={anchor.c}"
            )
        return LocatedNode(cell=cell)

    if isinstance(anchor, AnchorParagraph):
        paragraphs = _get_container_paragraphs(doc, anchor.container)
        if anchor.p >= len(paragraphs):
            raise AnchorNotFound(
                f"paragraph anchor 越界: p={anchor.p}, 容器仅 {len(paragraphs)} 段"
            )
        return LocatedNode(paragraph=paragraphs[anchor.p])

    if isinstance(anchor, AnchorRun):
        paragraphs = _get_container_paragraphs(doc, anchor.container)
        if anchor.p >= len(paragraphs):
            raise AnchorNotFound(
                f"run anchor paragraph 越界: p={anchor.p}, 容器仅 {len(paragraphs)} 段"
            )
        runs = paragraphs[anchor.p].runs
        if anchor.run_index >= len(runs):
            raise AnchorNotFound(
                f"run anchor run 越界: run_index={anchor.run_index}, 段仅 {len(runs)} run"
            )
        slice_range = None
        if anchor.char_offset is not None:
            slice_range = (anchor.char_offset, anchor.char_length or 0)
        return LocatedNode(run=runs[anchor.run_index], slice_range=slice_range)

    raise AnchorNotFound(f"未知 anchor 类型: {type(anchor)}")


# ====================== format_value ======================


_DATE_PATTERN_MAP = {
    "YYYY": "%Y",
    "MM": "%m",
    "DD": "%d",
    "HH": "%H",
    "mm": "%M",
    "ss": "%S",
}


def _convert_date_pattern(pattern: str) -> str:
    """把 'YYYY-MM-DD' 这种用户友好格式转成 strftime 格式。

    按 token 长度降序替换避免冲突（MM 不会先匹配 M）。
    """
    out = pattern
    for token in sorted(_DATE_PATTERN_MAP.keys(), key=len, reverse=True):
        out = out.replace(token, _DATE_PATTERN_MAP[token])
    return out


def _ensure_date(raw: Any) -> date:
    if isinstance(raw, datetime):
        return raw.date()
    if isinstance(raw, date):
        return raw
    if isinstance(raw, str):
        # 容错：支持 ISO 格式
        try:
            return datetime.fromisoformat(raw).date()
        except ValueError as e:
            raise ValueError(f"无法把 {raw!r} 解析为日期") from e
    raise ValueError(f"不支持的日期类型: {type(raw)}")


def format_value(raw: Any, fmt: Optional[str] = None) -> str:
    """把任意 raw 值按 fmt 格式化为字符串。

    支持的 fmt：
      - None        → str(raw)
      - "date:YYYY-MM-DD" / "date:YYYY年MM月DD日" / 等
      - "currency:CNY" / "currency:USD"
      - 不识别的 fmt → 退化为 str(raw)，不抛错（避免格式化阻塞渲染）

    raw 为 None → 返回空串。
    """
    if raw is None:
        return ""
    if fmt is None:
        return str(raw)

    if fmt.startswith("date:"):
        try:
            pattern = _convert_date_pattern(fmt.split(":", 1)[1])
            return _ensure_date(raw).strftime(pattern)
        except (ValueError, AttributeError):
            return str(raw)

    if fmt.startswith("currency:"):
        currency = fmt.split(":", 1)[1]
        try:
            amount = Decimal(str(raw))
            return f"{amount:,.2f} {currency}"
        except (ValueError, ArithmeticError):
            return str(raw)

    return str(raw)


# ====================== lxml rPr 深拷贝 ======================


def _extract_rpr_xml(paragraph: Paragraph) -> Optional[Any]:
    """从段落的第一个有 rPr 的 run 中深拷贝 rPr XML 元素。

    返回 lxml Element（已 deepcopy，可直接挂到新 run）或 None。
    """
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        if rpr is not None:
            return deepcopy(rpr)
    return None


def _apply_rpr_xml(run: Run, rpr_elem: Any) -> None:
    """把 rPr XML 元素挂到 run。会替换 run 现有的 rPr。"""
    r = run._r
    existing = r.find(qn("w:rPr"))
    if existing is not None:
        r.remove(existing)
    # rPr 必须是 run 的第一个子元素
    r.insert(0, rpr_elem)


def _clear_paragraph_runs(paragraph: Paragraph) -> None:
    """删除段落里所有 run 元素（不动 pPr）。"""
    p = paragraph._p
    for r in p.findall(qn("w:r")):
        p.remove(r)
    # 同时清掉 hyperlink 等会带 run 的元素
    for hl in p.findall(qn("w:hyperlink")):
        p.remove(hl)


def _clear_cell_keep_first_paragraph(cell: _Cell) -> None:
    """清空 cell 内容但保留第一个段落（cell 必须至少有一个段落）。"""
    tc = cell._tc
    paragraphs = tc.findall(qn("w:p"))
    # 保留第一个 <w:p>，删后面的
    for extra_p in paragraphs[1:]:
        tc.remove(extra_p)
    # 第一个段落清掉所有 run（保留 pPr）
    if paragraphs:
        first_p = paragraphs[0]
        for r in first_p.findall(qn("w:r")):
            first_p.remove(r)
        for hl in first_p.findall(qn("w:hyperlink")):
            first_p.remove(hl)


# ====================== apply_value ======================


def _write_to_cell(cell: _Cell, text: str) -> None:
    """cell anchor 的写入逻辑。"""
    # 判断 cell 是否"实质为空"（所有段落都无文字）
    has_content = any(
        (p.text or "").strip() for p in cell.paragraphs
    )

    if not has_content:
        # 干净路径：直接在第一段加一个 run，pPr/tcPr/边框/对齐全部不动
        para = cell.paragraphs[0]
        para.add_run(text)
        return

    # 脏路径：cell 已有内容（anchor 数据脏，保存时本应拒绝）
    # 降级：保留第一段，clone 第一个 run 的 rPr 给新 run
    template_rpr = _extract_rpr_xml(cell.paragraphs[0])
    _clear_cell_keep_first_paragraph(cell)
    para = cell.paragraphs[0]
    new_run = para.add_run(text)
    if template_rpr is not None:
        _apply_rpr_xml(new_run, template_rpr)


def _write_to_paragraph(paragraph: Paragraph, text: str) -> None:
    """paragraph anchor 的写入逻辑。沿用首 run rPr，pPr 不动。"""
    template_rpr = _extract_rpr_xml(paragraph)
    _clear_paragraph_runs(paragraph)
    new_run = paragraph.add_run(text)
    if template_rpr is not None:
        _apply_rpr_xml(new_run, template_rpr)


def _write_to_run(run: Run, text: str, slice_range: Optional[tuple[int, int]]) -> None:
    """run anchor 的写入逻辑。rPr 完全不动。"""
    if slice_range is None:
        run.text = text
        return
    offset, length = slice_range
    old = run.text or ""
    if offset + length > len(old):
        raise AnchorNotFound(
            f"run text too short: need offset+length={offset + length}, "
            f"actual len={len(old)} (text={old!r})"
        )
    run.text = old[:offset] + text + old[offset + length:]


def apply_value(
    doc: DocxDocument,
    anchor: BaseModel,
    value: Any,
    fmt: Optional[str] = None,
) -> None:
    """把 value（按 fmt 格式化后）写入 anchor 指向的位置。

    不变量：除目标 cell/paragraph/run 外，docx 其余 XML 字节级保持不变。
    失败抛 AnchorNotFound（路由层应捕获翻译成 HTTP 422）。
    """
    text = format_value(value, fmt)
    located = locate_anchor(doc, anchor)

    if isinstance(anchor, AnchorCell):
        assert located.cell is not None
        _write_to_cell(located.cell, text)
    elif isinstance(anchor, AnchorParagraph):
        assert located.paragraph is not None
        _write_to_paragraph(located.paragraph, text)
    elif isinstance(anchor, AnchorRun):
        assert located.run is not None
        _write_to_run(located.run, text, located.slice_range)
    else:
        raise AnchorNotFound(f"未知 anchor 类型: {type(anchor)}")
