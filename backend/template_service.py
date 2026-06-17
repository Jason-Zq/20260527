"""
Word 模板服务模块（v2 方案：可视化打标记 + LLM 语义匹配）。
- LLM 建议占位符（含 description）
- mammoth 将 docx 转 HTML 供前端预览
- 客户信息 LLM 语义匹配（缓存 > LLM 批量）
- docx 字符串替换渲染 + docx2pdf 转换
"""

import os
import re
import json
import asyncio
import base64
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, Tuple, List, Any

from docx import Document as DocxDocument
from docx.oxml.ns import qn as _docx_qn
import mammoth
from bs4 import BeautifulSoup
from sqlalchemy import select

import llm_service
from db.engine import async_session_maker
from db.models import Client, ClientInfo
from db import template_crud

# ---------- soffice 定位 + docx → PNG 渲染（高保真预览） ----------

_SOFFICE_CANDIDATE_PATHS = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/usr/bin/soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
]

_SOFFICE_CACHE: dict = {"path": None, "checked": False}

def _find_soffice() -> Optional[str]:
    """查找系统 LibreOffice soffice 可执行文件路径。"""
    if _SOFFICE_CACHE["checked"]:
        return _SOFFICE_CACHE["path"]

    # 1) PATH 查找
    for name in ("soffice", "soffice.exe", "libreoffice", "libreoffice.exe"):
        found = shutil.which(name)
        if found:
            _SOFFICE_CACHE["path"] = found
            break

    # 2) 常见安装路径
    if not _SOFFICE_CACHE["path"]:
        for cand in _SOFFICE_CANDIDATE_PATHS:
            if os.path.exists(cand):
                _SOFFICE_CACHE["path"] = cand
                break

    _SOFFICE_CACHE["checked"] = True
    if _SOFFICE_CACHE["path"]:
        print(f"[template_service] 找到 soffice: {_SOFFICE_CACHE['path']}")
    else:
        print(f"[template_service] 未找到 soffice，Word 原貌预览将不可用")
    return _SOFFICE_CACHE["path"]

def _docx_to_pdf_via_soffice(soffice: str, docx_path: str, pdf_path: str) -> None:
    """调用 soffice headless 把 docx 转成 pdf。失败抛 RuntimeError。"""
    pdf_dir = os.path.dirname(os.path.abspath(pdf_path)) or "."
    cmd = [
        soffice,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", pdf_dir,
        os.path.abspath(docx_path),
    ]
    try:
        result = subprocess.run(
            cmd,
            timeout=60,
            check=True,
            capture_output=True,
            text=True,
        )
    except subprocess.TimeoutExpired as e:
        raise RuntimeError(f"soffice 转换超时: {e}") from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"soffice 转换失败 (exit={e.returncode}): {e.stderr or e.stdout}"
        ) from e
    except FileNotFoundError as e:
        raise RuntimeError(f"soffice 可执行文件不存在: {soffice}") from e

    # soffice 默认输出文件名 = docx 的 basename + .pdf
    expected_pdf = os.path.join(pdf_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if not os.path.exists(expected_pdf):
        raise RuntimeError(f"soffice 未生成 PDF（期望路径: {expected_pdf}）")
    if os.path.normpath(expected_pdf) != os.path.normpath(pdf_path):
        shutil.move(expected_pdf, pdf_path)

def render_docx_pages(docx_path: str, output_dir: str, dpi: int = 150) -> List[str]:
    """
    docx → 多张 PNG 页面。
    返回 PNG 文件绝对路径列表（按页码顺序）。
    若 soffice 不可用抛 RuntimeError，让调用方决定降级。
    该函数应通过 asyncio.to_thread 调用。
    """
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError("未找到 soffice，无法生成 Word 原貌预览")

    os.makedirs(output_dir, exist_ok=True)

    # 1) docx → pdf（中间产物）
    with tempfile.TemporaryDirectory(prefix="docx_render_") as tmp:
        tmp_pdf = os.path.join(tmp, "intermediate.pdf")
        _docx_to_pdf_via_soffice(soffice, docx_path, tmp_pdf)

        # 2) pdf → png（用 pypdfium2，复用 ocr_service 模式）
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(tmp_pdf)
        page_paths: List[str] = []
        try:
            for i in range(len(pdf)):
                page = pdf[i]
                bitmap = page.render(scale=dpi / 72)
                pil_image = bitmap.to_pil()
                out_path = os.path.join(output_dir, f"page_{i + 1}.png")
                pil_image.save(out_path, "PNG")
                page_paths.append(os.path.abspath(out_path))
        finally:
            pdf.close()

    return page_paths

# ---------- mammoth 深度优化配置 ----------

_MAMMOTH_STYLE_MAP = """
p[style-name='Heading 1'] => h1:fresh
p[style-name='Heading 2'] => h2:fresh
p[style-name='Heading 3'] => h3:fresh
p[style-name='Heading 4'] => h4:fresh
p[style-name='Heading 5'] => h5:fresh
p[style-name='Heading 6'] => h6:fresh
p[style-name='Title'] => h1.doc-title:fresh
p[style-name='Subtitle'] => h2.doc-subtitle:fresh
p[style-name='Body Text'] => p.body-text:fresh
p[style-name='Normal'] => p.normal:fresh
p[style-name='List Paragraph'] => ul > li:fresh
p[style-name='List Bullet'] => ul > li:fresh
p[style-name='List Bullet 2'] => ul > li:fresh
p[style-name='List Bullet 3'] => ul > li:fresh
p[style-name='List Number'] => ol > li:fresh
p[style-name='List Number 2'] => ol > li:fresh
p[style-name='List Number 3'] => ol > li:fresh
p[style-name='Quote'] => blockquote:fresh
r[style-name='Strong'] => strong
r[style-name='Emphasis'] => em
table[style-name='Table Grid'] => table.table-grid:fresh
""".strip()

def _image_to_base64(image):
    """mammoth 图片转换器：把 docx 内嵌图片转成 data URL。"""
    try:
        with image.open() as img_bytes:
            encoded = base64.b64encode(img_bytes.read()).decode("utf-8")
        content_type = image.content_type or "image/png"
        return {"src": f"data:{content_type};base64,{encoded}"}
    except Exception as e:
        print(f"[template_service] 图片 base64 失败: {e}")
        return {"src": ""}

def _has_border(tc) -> bool:
    """判断 docx 单元格是否有边框定义。"""
    tc_pr = tc.find(_docx_qn("w:tcPr"))
    if tc_pr is None:
        return False
    tc_borders = tc_pr.find(_docx_qn("w:tcBorders"))
    if tc_borders is None:
        return False
    for side in ("top", "left", "bottom", "right"):
        el = tc_borders.find(_docx_qn(f"w:{side}"))
        if el is not None:
            val = (el.get(_docx_qn("w:val")) or "").lower()
            if val and val not in ("none", "nil"):
                return True
    return False

def _post_process_tables(html: str, docx_path: str) -> str:
    """
    mammoth 表格输出往往丢失边框和合并单元格信息。
    本函数用 python-docx 读取原始表格结构，对 mammoth HTML 做后处理：
    - 为无边框表格加 class="table-with-borders"
    - 为单元格注入 colspan / rowspan（基于 gridSpan / vMerge）
    """
    if not html or "<table" not in html:
        return html
    try:
        doc = DocxDocument(docx_path)
    except Exception as e:
        print(f"[template_service] docx 解析失败，跳过后处理: {e}")
        return html

    soup = BeautifulSoup(html, "html.parser")
    html_tables = soup.find_all("table")
    docx_tables = doc.tables
    n = min(len(html_tables), len(docx_tables))

    for t_idx in range(n):
        html_table = html_tables[t_idx]
        docx_table = docx_tables[t_idx]

        # 统计是否有单元格带边框，若全部无边框则给整表加 class
        has_any_border = False
        for row in docx_table.rows:
            for cell in row.cells:
                if _has_border(cell._tc):
                    has_any_border = True
                    break
            if has_any_border:
                break

        if not has_any_border:
            existing_cls = html_table.get("class") or []
            if "table-with-borders" not in existing_cls:
                html_table["class"] = existing_cls + ["table-with-borders"]

        # 合并单元格映射：(row_idx, col_idx) -> {colspan, rowspan}
        merge_map: dict[tuple[int, int], dict] = {}
        for row_idx, row in enumerate(docx_table.rows):
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tc_pr = tc.find(_docx_qn("w:tcPr"))
                if tc_pr is None:
                    continue
                # 水平合并 gridSpan
                gs = tc_pr.find(_docx_qn("w:gridSpan"))
                if gs is not None:
                    try:
                        span = int(gs.get(_docx_qn("w:val")) or "1")
                        if span > 1:
                            merge_map.setdefault((row_idx, col_idx), {})["colspan"] = span
                    except ValueError:
                        pass
                # 垂直合并 vMerge（continue 表示当前单元格是被合并的，restart 表示新合并起点）
                vm = tc_pr.find(_docx_qn("w:vMerge"))
                if vm is not None:
                    val = vm.get(_docx_qn("w:val"), "continue")
                    if val == "restart":
                        height = 1
                        for r in docx_table.rows[row_idx + 1:]:
                            try:
                                next_tc = r.cells[col_idx]._tc
                            except Exception:
                                break
                            npr = next_tc.find(_docx_qn("w:tcPr"))
                            nvm = npr.find(_docx_qn("w:vMerge")) if npr is not None else None
                            if nvm is None:
                                break
                            if nvm.get(_docx_qn("w:val"), "continue") == "restart":
                                break
                            height += 1
                        if height > 1:
                            merge_map.setdefault((row_idx, col_idx), {})["rowspan"] = height

        if not merge_map:
            continue

        # 把合并信息注入 HTML <td>
        html_rows = html_table.find_all("tr", recursive=False)
        # 兼容 tbody 包裹
        if not html_rows:
            tbody = html_table.find("tbody")
            if tbody:
                html_rows = tbody.find_all("tr", recursive=False)
        for row_idx, html_tr in enumerate(html_rows):
            tds = html_tr.find_all(["td", "th"], recursive=False)
            for col_idx, html_td in enumerate(tds):
                info = merge_map.get((row_idx, col_idx))
                if not info:
                    continue
                if "colspan" in info:
                    html_td["colspan"] = str(info["colspan"])
                if "rowspan" in info:
                    html_td["rowspan"] = str(info["rowspan"])

    return str(soup)

# ---------- 文档展示 ----------

def convert_docx_to_html(docx_path: str) -> str:
    """mammoth 将 docx 转 HTML，并做深度优化（style_map / 图片 base64 / 表格后处理）。"""
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"docx 文件不存在: {docx_path}")
    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(
            f,
            style_map=_MAMMOTH_STYLE_MAP,
            include_default_style_map=True,
            convert_image=mammoth.images.img_element(_image_to_base64),
        )
    html = result.value or ""
    html = _post_process_tables(html, docx_path)
    return html

# ---------- 结构化扫描：anchor 候选 ----------

_UNDERSCORE_RE = re.compile(r"_{3,}")

def _find_label_for_cell(table, r_idx: int, c_idx: int) -> str:
    """
    推断单元格的语义 label：
      1) 同行左边最近的非空 cell 文本
      2) 上一行同列的非空 cell 文本
    找不到返回空字符串。
    """
    # 同行左边（跳过被合并的虚拟 cell：python-docx 在 row.cells 中重复返回合并的 cell）
    try:
        row = table.rows[r_idx]
        for prev_c in range(c_idx - 1, -1, -1):
            t = (row.cells[prev_c].text or "").strip()
            if t:
                return t
    except (IndexError, AttributeError):
        pass
    # 上行同列
    if r_idx > 0:
        try:
            t = (table.rows[r_idx - 1].cells[c_idx].text or "").strip()
            if t:
                return t
        except (IndexError, AttributeError):
            pass
    return ""

def _extract_label_before(text_before: str) -> str:
    """
    从 'X：' / 'X:' / 'X（描述）：' 模式中提取 X 作为 label。
    输入是从下划线起始位置往前的内容（不含下划线本身）。

    步骤：
      1) 去尾随空白 + 冒号
      2) 移除括号内补充说明（如 'X（跨 4 列）' → 'X'）
      3) 再去冒号
      4) 取末尾连续非标点段作为 label
    """
    if not text_before:
        return ""
    s = text_before.rstrip()
    s = s.rstrip(":：").rstrip()
    if not s:
        return ""
    # 移除括号内补充说明（中文/英文括号）
    s = re.sub(r"[（(][^）)]*[)）]", "", s)
    s = s.rstrip(":：").rstrip()
    if not s or re.fullmatch(r"[_\s_:：:、()（）]+", s):
        return ""
    # 取末尾连续非标点段
    m = re.search(r"([^\s_:：:、()（）]+)\s*$", s)
    if not m:
        return ""
    label = m.group(1).strip()
    if not label or re.fullmatch(r"[_：:、()（）\s]+", label):
        return ""
    return label

# ---------- v2 Anchor 扫描器 ----------

def _scan_runs_for_underscore_anchors(
    runs, container: dict, p_idx: int, fallback_label: str, out: list[dict]
) -> None:
    """在一组 run 里找 `___+` 模式，每处生成一个 run anchor。

    fallback_label：当下划线左侧文本提不出 label 时用的备选。
    """
    # 先把整段文字拼起来，便于按"下划线左侧文字"提取更精准的 label
    full_text_parts: list[tuple[int, str]] = []  # [(run_idx, text)]
    cursor = 0
    run_offsets: list[tuple[int, int]] = []  # [(start_in_full, end_in_full)] per run
    for run in runs:
        text = run.text or ""
        run_offsets.append((cursor, cursor + len(text)))
        full_text_parts.append((len(run_offsets) - 1, text))
        cursor += len(text)
    full_text = "".join(text for _, text in full_text_parts)

    for run_idx, run in enumerate(runs):
        text = run.text or ""
        for m in _UNDERSCORE_RE.finditer(text):
            # label：先看下划线在 full_text 中位置之前的部分能否提取
            run_start, _ = run_offsets[run_idx]
            full_start = run_start + m.start()
            label_from_text = _extract_label_before(full_text[:full_start].rstrip())
            label = label_from_text or fallback_label

            out.append({
                "anchor": {
                    "kind": "run",
                    "container": container,
                    "p": p_idx,
                    "run_index": run_idx,
                    "char_offset": m.start(),
                    "char_length": m.end() - m.start(),
                },
                "label_context": label,
            })

def scan_anchors(docx_path: str) -> list[dict]:
    """v2 扫描器：找出 docx 里所有需填写的位置，输出 anchor 候选列表。

    返回 `[{anchor: dict, label_context: str}, ...]`，其中 anchor 与
    [backend/anchor.py](backend/anchor.py) 的 Pydantic 模型字段一一对应，
    可直接 JSON 序列化、可直接喂给 `apply_value`。

    覆盖的占位形态：
      - 表格空 cell → AnchorCell
      - 表格 cell 内单 run 中的 `___+` 下划线 → AnchorRun(char_offset/length)
      - body 段落内单 run 中的 `___+` 下划线 → AnchorRun(char_offset/length)

    不覆盖（v1 限制）：
      - 跨 run 的下划线（同一段落被 Word 切成多 run）—— 极少见，用户可手动添加
      - 段落里非下划线的占位文字（如 "X 年 X 月" 中的 X）
      - header / footer（业务场景少）

    不调用 LLM，纯结构扫描。
    """
    from docx import Document as _Docx
    doc = _Docx(docx_path)
    out: list[dict] = []

    # 1) 表格
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            seen_tc_ids: set = set()
            for c_idx, cell in enumerate(row.cells):
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue
                seen_tc_ids.add(tc_id)

                text = cell.text or ""
                stripped = text.strip()

                if not stripped:
                    # 空 cell → AnchorCell
                    label = _find_label_for_cell(table, r_idx, c_idx)
                    out.append({
                        "anchor": {"kind": "cell", "t": t_idx, "r": r_idx, "c": c_idx},
                        "label_context": label,
                    })
                else:
                    # cell 非空：扫 cell 内每个段落、每个 run 找下划线
                    cell_label = _find_label_for_cell(table, r_idx, c_idx)
                    container = {"kind": "cell", "t": t_idx, "r": r_idx, "c": c_idx}
                    for p_idx, para in enumerate(cell.paragraphs):
                        _scan_runs_for_underscore_anchors(
                            para.runs, container, p_idx, cell_label, out
                        )

    # 2) body 段落（python-docx 的 doc.paragraphs 不含表格内段落）
    body_container = {"kind": "body"}
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        if not text:
            continue
        _scan_runs_for_underscore_anchors(para.runs, body_container, p_idx, "", out)

    return out

# ---------- v2 LLM 描述增强 + 客户匹配（field_hint 规则优先） ----------

def enrich_anchors_with_llm(anchors: list[dict], docx_text: str | None = None) -> list[dict]:
    """为每个 anchor 生成 description + field_hint。

    输入：scan_anchors 输出的 `[{anchor, label_context}, ...]` 列表
    输出：每个 anchor 加上 `description` 和 `field_hint`（key 用 FIELD_DICTIONARY.key）

    策略：
      - 先用 FIELD_DICTIONARY 文本匹配把能识别的字段直接归类（field_hint + description）
      - 剩下的"未识别"anchor 一次性送 LLM 批量分类
      - description 取 label_context 优先，LLM 没给出时回退 label_context

    docx_text 用于喂 LLM 的上下文（可选，省略时不传）。
    """
    import field_dictionary as fd

    enriched: list[dict] = []
    need_llm: list[dict] = []  # idx_in_enriched 给 LLM 用

    for a in anchors:
        new = dict(a)  # shallow copy
        label = a.get("label_context") or ""
        # 规则层：label 命中字典
        match = fd.find_field_by_text(label)
        if match:
            new["field_hint"] = match.key
            new["description"] = label  # 用 scan 拿到的 label 当 description
            new["default_fmt"] = match.fmt
        else:
            # 等 LLM 给
            new["field_hint"] = None
            new["description"] = label
            new["default_fmt"] = None
            need_llm.append(new)
        enriched.append(new)

    if not need_llm:
        return enriched

    # 调 LLM：给候选 field_hint 列表 + 每个未识别 anchor 的 label
    candidates_for_llm = [
        {"idx": enriched.index(a), "label_context": a.get("label_context") or ""}
        for a in need_llm
    ]
    prompt = (
        "你是表单字段分类助手。下面的每个「anchor」是一个 Word 模板里的待填位置，"
        "我已经从结构上拿到了它的 label_context（来自占位符左侧文字/单元格邻居）。\n"
        "请你为每个 anchor 在下面【候选字段列表】中挑一个最匹配的 key；"
        "如果完全找不到合适匹配，返回空字符串。\n\n"
        "【候选字段列表】\n"
        f"{json.dumps(fd.list_for_llm_prompt(), ensure_ascii=False, indent=2)}\n\n"
        "【anchor 列表】\n"
        f"{json.dumps(candidates_for_llm, ensure_ascii=False, indent=2)}\n\n"
        '返回 JSON：{"results": [{"idx": <number>, "field_hint": "<key or empty>"}]}\n'
        "只返回 JSON，不要其他文字。"
    )
    if docx_text:
        prompt += f"\n\n参考文档上下文：\n{docx_text[:3000]}"

    try:
        from llm_service import _call_llm
        result_text = _call_llm(prompt)
        data = json.loads(result_text)
        results = data.get("results") or []
        for r in results:
            idx = r.get("idx")
            hint = (r.get("field_hint") or "").strip()
            if not isinstance(idx, int):
                continue
            if not (0 <= idx < len(enriched)):
                continue
            if hint and hint in [fd.key for fd in fd.FIELD_DICTIONARY]:
                match = fd.get_field(hint)
                enriched[idx]["field_hint"] = hint
                enriched[idx]["default_fmt"] = match.fmt if match else None
    except Exception as e:
        print(f"[template_service] enrich LLM 失败（保留规则层结果）: {e}")

    return enriched

async def _load_client_sources_v2(client_id: int) -> dict:
    """v2 客户字段拍平：主表 + client_info + computed。

    返回 {key: value}，key 用 FIELD_DICTIONARY 的 key。
    """
    import field_dictionary as fd
    from datetime import datetime, date

    async with async_session_maker() as session:
        res = await session.execute(select(Client).where(Client.id == client_id))
        client = res.scalar_one_or_none()
        if not client:
            return {}

        sources: dict = {}
        # 主表 6 字段
        if client.name:
            sources["name"] = client.name
        if client.id_number:
            sources["id_number"] = client.id_number
        if client.gender:
            sources["gender"] = client.gender
        if client.birth_date:
            sources["birth_date"] = client.birth_date
        if client.nationality:
            sources["nationality"] = client.nationality
        if client.consultant:
            sources["consultant"] = client.consultant
        if client.notes:
            sources["notes"] = client.notes

        # client_info KV → 按字典的 source 反查匹配 key
        info_res = await session.execute(
            select(ClientInfo).where(ClientInfo.client_id == client_id)
        )
        infos = info_res.scalars().all()

        # 建反向索引：info_key → field key
        info_key_to_fd_key: dict[str, str] = {}
        for fdef in fd.FIELD_DICTIONARY:
            if fdef.source.startswith("client_info."):
                ik = fdef.source.split(".", 1)[1]
                info_key_to_fd_key[ik] = fdef.key

        for info in infos:
            if not info.info_value:
                continue
            # 优先按字典精确匹配
            fd_key = info_key_to_fd_key.get(info.info_key)
            if fd_key:
                sources[fd_key] = info.info_value
            else:
                # 兜底：用 info_key 作 key（业务自定义字段）
                sources[info.info_key] = info.info_value

        # computed 字段
        for fdef in fd.FIELD_DICTIONARY:
            if fdef.source == "computed.today":
                sources["today"] = date.today()

        return sources

async def match_anchors_to_client(
    anchors: list[dict],
    client_id: int,
    template_id: int | None = None,
) -> dict:
    """按 anchor 列表做客户字段匹配。

    返回 {"matched": {placeholder_id: value}, "unmatched": [placeholder_id], "from_cache": bool}

    策略：
      1) 缓存（template_id + client_id）命中即返回
      2) 规则层：anchor 有 field_hint → 从 sources 直查，命中即返回
      3) LLM 兜底层：剩下没匹配或没 hint 的 anchor 批量送 LLM
    """
    import field_dictionary as fd

    if not anchors:
        return {"matched": {}, "unmatched": [], "from_cache": False}

    # placeholder id = strN
    ids = [f"str{i+1}" for i in range(len(anchors))]

    # 1) 缓存
    if template_id is not None:
        cached = await template_crud.get_cached_fill(template_id, client_id)
        if cached:
            matched = {i: str(cached[i]) for i in ids if i in cached and cached[i]}
            unmatched = [i for i in ids if i not in matched]
            return {"matched": matched, "unmatched": unmatched, "from_cache": True}

    # 2) 规则层
    sources = await _load_client_sources_v2(client_id)
    matched: dict[str, str] = {}
    need_llm: list[int] = []  # idx_in_anchors 待送 LLM

    for idx, ph in enumerate(anchors):
        ph_id = ids[idx]
        hint = ph.get("field_hint")
        if hint and hint in sources:
            matched[ph_id] = str(sources[hint])
        else:
            need_llm.append(idx)

    if not need_llm:
        unmatched = [ids[i] for i in range(len(anchors)) if ids[i] not in matched]
        return {"matched": matched, "unmatched": unmatched, "from_cache": False}

    # 3) LLM 兜底
    llm_payload = []
    for idx in need_llm:
        llm_payload.append({
            "placeholder_id": ids[idx],
            "description": anchors[idx].get("description") or "",
            "label_context": anchors[idx].get("label_context") or "",
            "candidate_keys": [f.key for f in fd.FIELD_DICTIONARY],
        })

    sources_view = {k: str(v) for k, v in sources.items() if not hasattr(v, 'isoformat') or True}
    # 把 date/date 序列化（str 化）
    for k, v in list(sources_view.items()):
        if hasattr(v, 'isoformat'):
            sources_view[k] = v.isoformat()

    prompt = (
        "你是字段匹配助手。下面是模板占位符（含描述），请从「客户已有信息」中找最合适的值。\n"
        "如果客户信息里没有合适匹配，对应占位符返回空字符串。\n"
        "只返回 JSON，键为 placeholder_id，值为对应客户信息值。\n\n"
        f"占位符列表：\n{json.dumps(llm_payload, ensure_ascii=False, indent=2)}\n\n"
        f"客户已有信息：\n{json.dumps(sources_view, ensure_ascii=False, indent=2)}\n\n"
        '示例：{"str1": "张三", "str2": "E1234567"}'
    )

    try:
        from llm_service import _call_llm
        result_text = await asyncio.to_thread(_call_llm, prompt)
        data = json.loads(result_text)
        if isinstance(data, dict):
            for ph_id, v in data.items():
                if ph_id in ids and v and str(v).strip():
                    matched[ph_id] = str(v).strip()
    except Exception as e:
        print(f"[template_service] match LLM 兜底失败: {e}")

    unmatched = [i for i in ids if i not in matched]
    return {"matched": matched, "unmatched": unmatched, "from_cache": False}

# ---------- 客户信息映射（v2：LLM 语义匹配主路径） ----------

# ---------- PDF 渲染 ----------

def _convert_docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    """
    docx2pdf 走 Word COM；FastAPI 异步上下文需要 CoInitialize。
    本函数应在 asyncio.to_thread 中调用。
    """
    import pythoncom
    from docx2pdf import convert as docx2pdf_convert

    pythoncom.CoInitialize()
    try:
        docx2pdf_convert(docx_path, pdf_path)
    finally:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass

# ---------- v2 渲染：apply_value 驱动的批量渲染 ----------

def _anchor_sort_key(item: tuple[dict, str]) -> tuple:
    """batch_apply_anchors 的排序 key：同 run 上 offset 降序；不同 run 任意。

    item: (anchor_dict, value)
    返回 (anchor_kind_rank, container_path, run_index_neg_offset)

    排序原则：
      - cell anchor 排最前（互不影响）
      - paragraph anchor 次之
      - run anchor 排最后；同 (container, p, run_index) 内 offset 降序
    """
    a = item[0]
    kind = a.get("kind")
    if kind == "cell":
        return (0, (a["t"], a["r"], a["c"]), 0)
    if kind == "paragraph":
        container = a.get("container") or {}
        return (1, _container_key(container), 0)
    # run
    container = a.get("container") or {}
    run_index = a.get("run_index", 0)
    char_offset = a.get("char_offset") or 0
    # 用 (1 - offset) 当第三位：offset 大的排前面
    return (2, _container_key(container) + (run_index,), -char_offset)

def _container_key(container: dict) -> tuple:
    """把 container dict 转成可比较的 tuple。"""
    kind = container.get("kind")
    if kind == "body":
        return ("body",)
    if kind == "cell":
        return ("cell", container.get("t"), container.get("r"), container.get("c"))
    if kind == "header":
        return ("header", container.get("s"))
    if kind == "footer":
        return ("footer", container.get("s"))
    return ("unknown",)

def batch_apply_anchors(
    docx_path: str,
    anchor_items: list[tuple[dict, Any, Optional[str]]],
    out_path: str,
) -> None:
    """批量应用 anchor 值到 docx，写到 out_path。

    anchor_items: [(anchor_dict, value, fmt_or_None), ...]
    按 cell → paragraph → run 排序；run 内同 (container,p,run_index) 的按 offset 降序，
    保证同 run 多次 apply 时前面的 anchor 不被后面的破坏。

    应通过 asyncio.to_thread 调用（docx 操作同步）。
    """
    from docx import Document as _Docx

    # 排序
    sorted_items = sorted(anchor_items, key=_anchor_sort_key)

    doc = _Docx(docx_path)
    from anchor import AnchorCell, AnchorParagraph, AnchorRun, apply_value

    for anchor_dict, value, fmt in sorted_items:
        if not anchor_dict:
            continue
        kind = anchor_dict.get("kind")
        if kind == "cell":
            anchor = AnchorCell(**anchor_dict)
        elif kind == "paragraph":
            anchor = AnchorParagraph(**anchor_dict)
        elif kind == "run":
            anchor = AnchorRun(**anchor_dict)
        else:
            continue
        if value is None or value == "":
            continue  # 跳过空值
        try:
            apply_value(doc, anchor, value, fmt=fmt)
        except Exception as e:
            # 跳过失效 anchor，记录日志不中断
            print(f"[template_service] batch_apply 跳过 anchor: {anchor_dict} → {e}")
            continue

    doc.save(out_path)

def render_preview_html_v2(
    anchors: list[dict],
    anchor_values: dict[str, Any],
    base_html: str,
) -> str:
    """v2 预览 HTML：在 mammoth 转出的 HTML 上按 anchor.kind 不同方式高亮占位符。

    与 v1 的差别：v1 靠 original_text 字符串替换；v2 靠 mammoth 转出的元素位置。
    mammoth 把表格 cell/段落/run 都尽可能保留，因此 cell anchor 可以在 cell 文本里
    替换标记字符串。我们用同样的方式：
      - cell anchor：在 cell 文本里插一个可见的 strN 标记
      - run anchor：在 run 文本的 char_offset 范围插一个 strN 标记
      - paragraph anchor：整段替换为 strN 标记

    简单起见，v2 预览不 1:1 还原物理位置，只把"已填值 / 未填占位符"标出来。
    """
    from html import escape as _esc
    if not base_html:
        return ""

    # 把每个 anchor 转成"（值或 strN）"文本
    snippet_map = {}
    for ph in anchors:
        ph_id = ph.get("id") or ""
        v = anchor_values.get(ph_id)
        if v:
            snippet_map[ph_id] = (
                f'<span class="pv-filled" data-id="{_esc(ph_id)}">'
                f'{_esc(str(v))}</span>'
            )
        else:
            snippet_map[ph_id] = (
                f'<span class="pv-empty" data-id="{_esc(ph_id)}">{_esc(ph_id)}</span>'
            )

    # 对 cell anchor：找到 cell 文本里含 "strN" 的位置（mammoth 不直接给 cell 锚点）
    # 这里用纯 HTML 文本扫描：找到第一个空 cell 文本节点 → 替换为 snippet
    # 但 mammoth 转出的 cell 可能没特殊标记，难度高
    # 简化方案：v2 预览返回 base_html，前端用 anchor + 预览数据自己画 overlay
    return base_html

async def render_to_pdf_v2(
    template_id: int,
    anchor_values: dict[str, Any],
    fmt_overrides: dict[str, str] | None = None,
) -> Tuple[Optional[str], Optional[str]]:
    """v2 渲染：基于 anchor 的 apply_value 渲染。

    返回 (pdf_path, fallback_docx_path)。
    """
    tpl_dict = await template_crud.get_template_dict(template_id)
    if not tpl_dict:
        raise ValueError(f"模板 {template_id} 不存在")

    docx_template_path = tpl_dict["file_path"]
    if not docx_template_path or not os.path.isabs(docx_template_path):
        backend_dir = os.path.dirname(os.path.abspath(__file__))
        docx_template_path = os.path.normpath(os.path.join(backend_dir, "..", docx_template_path or ""))
    if not os.path.exists(docx_template_path):
        raise FileNotFoundError(f"模板文件不存在: {docx_template_path}")

    backend_dir = os.path.dirname(os.path.abspath(__file__))
    fills_dir = os.path.normpath(
        os.path.join(backend_dir, "..", "output", "templates", str(template_id), "fills")
    )
    os.makedirs(fills_dir, exist_ok=True)

    from datetime import datetime
    stamp = datetime.now().strftime("%y%m%d%H%M%S")
    docx_out = os.path.join(fills_dir, f"{stamp}.docx")
    pdf_out = os.path.join(fills_dir, f"{stamp}.pdf")

    anchors = tpl_dict.get("placeholders") or []

    # 组装 anchor_items: (anchor_dict, value, fmt)
    fmt_overrides = fmt_overrides or {}
    items: list[tuple[dict, Any, Optional[str]]] = []
    for ph in anchors:
        if not isinstance(ph, dict):
            continue
        ph_id = ph.get("id")
        anchor = ph.get("anchor")
        if not ph_id or not anchor:
            continue
        value = anchor_values.get(ph_id, "")
        fmt = fmt_overrides.get(ph_id) or ph.get("default_fmt")
        items.append((anchor, value, fmt))

    def _render():
        batch_apply_anchors(docx_template_path, items, docx_out)

    await asyncio.to_thread(_render)

    try:
        await asyncio.to_thread(_convert_docx_to_pdf, docx_out, pdf_out)
        if os.path.exists(pdf_out):
            return pdf_out, None
        return None, docx_out
    except Exception as e:
        print(f"[template_service] PDF 转换失败，降级为 docx：{e}")
        return None, docx_out
