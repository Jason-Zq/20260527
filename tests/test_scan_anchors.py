"""
scan_anchors 单测 —— 验证扫描器输出的 anchor 都能被 locate_anchor 找回，
且能用 apply_value 正确填值。

运行：
    cd e:/qoderproject/20260527
    PYTHONIOENCODING=utf-8 ./.venv312/Scripts/python.exe tests/test_scan_anchors.py
"""

from __future__ import annotations

import io
import os
import sys
import traceback

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx import Document
from pydantic import TypeAdapter

from anchor import (
    Anchor,
    AnchorCell,
    AnchorRun,
    ContainerBody,
    ContainerCell,
    apply_value,
    locate_anchor,
)
from template_service import scan_anchors


# Pydantic adapter：把 scan_anchors 输出的 dict 解析回模型
_ANCHOR_ADAPTER = TypeAdapter(Anchor)


def _parse_anchor(raw: dict):
    """把 scan_anchors 返回的 anchor dict 转成 Pydantic 模型实例。"""
    return _ANCHOR_ADAPTER.validate_python(raw)


# ====================== 测试用 docx 构造 ======================


def make_underscore_docx():
    """构造一份含多种下划线占位的 docx：
      - body 段落：'姓名：______'（同段两个独立下划线）
      - body 段落：'国籍：__________'
      - 表格 cell 内：'签发机关：__________'（cell 非空）
      - 表格空 cell（无内容）
    """
    doc = Document()
    # 段落 1
    p1 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p1.add_run("姓名：______ 性别：____")  # 注意有两个 _ 段
    # 段落 2
    p2 = doc.add_paragraph()
    p2.add_run("国籍：__________")
    # 表格
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "签发机关"
    table.rows[0].cells[1].text = "签发日期：__________"
    table.rows[1].cells[0].text = "有效期"
    # rows[1].cells[1] 留空 → 空 cell
    return doc


def save_to_tmp(doc) -> str:
    """save docx 到临时路径，返回路径。"""
    import tempfile
    fd, path = tempfile.mkstemp(suffix=".docx", prefix="scan_test_")
    os.close(fd)
    doc.save(path)
    return path


# ====================== 测试用例 ======================


def test_01_underscore_scan_basic():
    """基础：构造的下划线文档应扫出 5 个 anchor（2 段落下划线 + 1 cell 内下划线 + 1 空 cell + 注意：第一段两个 ____ 分别算）。"""
    doc = make_underscore_docx()
    path = save_to_tmp(doc)
    try:
        results = scan_anchors(path)

        # 分类统计
        cell_anchors = [r for r in results if r["anchor"]["kind"] == "cell"]
        run_anchors = [r for r in results if r["anchor"]["kind"] == "run"]

        # 期望：1 个空 cell（rows[1].cells[1]）
        assert len(cell_anchors) == 1, f"空 cell 期望 1 实际 {len(cell_anchors)}"
        # 期望：4 个 run anchor（段1的 2 个下划线 + 段2的 1 个 + cell 内 1 个）
        assert len(run_anchors) == 4, f"下划线 anchor 期望 4 实际 {len(run_anchors)}"

        # 每个 anchor 都能被 locate_anchor 找回
        doc2 = Document(path)
        for r in results:
            anchor = _parse_anchor(r["anchor"])
            located = locate_anchor(doc2, anchor)
            if anchor.kind == "cell":
                assert located.cell is not None
            else:
                assert located.run is not None
    finally:
        os.remove(path)


def test_02_anchor_label_context_populated():
    """label_context 应从下划线左侧 'X：' 模式提取。"""
    doc = make_underscore_docx()
    path = save_to_tmp(doc)
    try:
        results = scan_anchors(path)
        labels = {r["label_context"] for r in results}
        # 应包含从段落/cell 文字推断出的关键 label
        # '姓名' / '性别' / '国籍' / '签发日期'（cell 内右侧的下划线）
        # 空 cell 应取上方/左侧 label：'有效期' 或 '签发日期'（左侧）
        expected_subset = {"姓名", "性别", "国籍"}
        missing = expected_subset - labels
        assert not missing, f"label_context 缺失: {missing}, 实际 {labels}"
    finally:
        os.remove(path)


def test_03_apply_value_on_scanned_anchors():
    """对每个扫出的 anchor 调 apply_value，应能正确写值不抛错。

    注意：每个 anchor 独立 reload docx，避免同 run 上多 anchor 互相影响偏移。
    这是 anchor 的设计选择——batch apply 同一 run 时需 offset 降序 + 长度对齐，
    阶段 3 的批量渲染器负责处理。阶段 2 只验证单 anchor 正交性。
    """
    doc = make_underscore_docx()
    path = save_to_tmp(doc)
    try:
        results = scan_anchors(path)
        assert len(results) >= 5, f"scan 期望 >=5 anchor，实际 {len(results)}"

        for i, r in enumerate(results):
            # 每个 anchor 独立 reload
            doc_fresh = Document(path)
            anchor = _parse_anchor(r["anchor"])
            apply_value(doc_fresh, anchor, f"V{i}")

            # 序列化重读验证
            buf = io.BytesIO()
            doc_fresh.save(buf)
            buf.seek(0)
            doc_round = Document(buf)
            located = locate_anchor(doc_round, anchor)
            if anchor.kind == "cell":
                assert located.cell.text == f"V{i}", (
                    f"anchor #{i} 写后验证失败: {located.cell.text!r}"
                )
            else:
                # run 写后：可能被替换的位置就是 char_offset 范围
                actual = located.run.text
                # 不强断言精确文本（不同 anchor 形态下新值插入位置有差异）
                # 关键是 V{i} 出现
                assert f"V{i}" in actual, (
                    f"anchor #{i} (run) 写后验证失败: {actual!r}"
                )
    finally:
        os.remove(path)


def test_04_poa_integration():
    """POA 信息表.docx 集成：4 表全是空 cell，应扫出大量 cell anchor。"""
    src = os.path.join(os.path.dirname(__file__), "..", "POA 信息表.docx")
    if not os.path.exists(src):
        raise AssertionError(f"POA 信息表.docx 不存在: {src}")

    results = scan_anchors(src)
    cell_anchors = [r for r in results if r["anchor"]["kind"] == "cell"]
    run_anchors = [r for r in results if r["anchor"]["kind"] == "run"]

    # POA 每表 12-13 行，含 1 行表头 + ~9 行"标签|空 cell"
    # 4 张表 → 期望 cell anchor 在 [30, 50] 区间
    assert 30 <= len(cell_anchors) <= 60, (
        f"POA cell anchor 期望 30-60 实际 {len(cell_anchors)}"
    )
    print(f"  POA scan: cell={len(cell_anchors)}, run={len(run_anchors)}")

    # 至少应有标准 5 字段（姓名 / 性别 / 出生日期 / 护照号 / 关系）的 label
    labels = {r["label_context"] for r in cell_anchors}
    # POA 模板第 1 列是中英混合，如 'Name/姓名'
    has_name = any("姓名" in lab or "Name" in lab for lab in labels)
    has_passport = any("护照" in lab or "Passport" in lab for lab in labels)
    assert has_name, f"应扫出姓名 label, 实际 labels: {sorted(labels)[:5]}"
    assert has_passport, f"应扫出护照号 label, 实际 labels: {sorted(labels)[:5]}"

    # 所有 cell anchor 都能 locate
    doc = Document(src)
    for r in cell_anchors:
        anchor = _parse_anchor(r["anchor"])
        located = locate_anchor(doc, anchor)
        assert located.cell is not None


def test_05_poa_fill_first_table():
    """POA 第 1 张表的姓名/性别/出生日期填上后，能 round-trip 序列化。"""
    src = os.path.join(os.path.dirname(__file__), "..", "POA 信息表.docx")
    if not os.path.exists(src):
        raise AssertionError(f"POA 信息表.docx 不存在: {src}")

    results = scan_anchors(src)
    # 取第 1 张表（t=0）的前 5 个空 cell anchor
    t0_cells = [
        _parse_anchor(r["anchor"])
        for r in results
        if r["anchor"]["kind"] == "cell" and r["anchor"]["t"] == 0
    ][:5]
    assert len(t0_cells) == 5, f"第 1 张表前 5 个 cell anchor 不足: {len(t0_cells)}"

    doc = Document(src)
    fills = ["张三", "妻子", "男", "1990-01-01", "E12345678"]
    for anchor, val in zip(t0_cells, fills):
        apply_value(doc, anchor, val)

    # 序列化重读
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    # 验证写入成功
    for anchor, val in zip(t0_cells, fills):
        located = locate_anchor(doc2, anchor)
        assert located.cell.text == val, (
            f"anchor t={anchor.t} r={anchor.r} c={anchor.c} 期望 {val!r} "
            f"实际 {located.cell.text!r}"
        )


# ====================== Runner ======================


TESTS = [
    ("01_underscore_scan_basic",              test_01_underscore_scan_basic),
    ("02_anchor_label_context_populated",     test_02_anchor_label_context_populated),
    ("03_apply_value_on_scanned_anchors",     test_03_apply_value_on_scanned_anchors),
    ("04_poa_integration",                    test_04_poa_integration),
    ("05_poa_fill_first_table",               test_05_poa_fill_first_table),
]


def main():
    passed = 0
    failed = []
    for name, fn in TESTS:
        try:
            fn()
            print(f"  PASS  {name}")
            passed += 1
        except Exception as e:
            print(f"  FAIL  {name}: {type(e).__name__}: {e}")
            traceback.print_exc()
            failed.append(name)
    total = len(TESTS)
    print(f"\n{passed}/{total} 通过")
    if failed:
        print("失败:", ", ".join(failed))
        sys.exit(1)
    sys.exit(0)


if __name__ == "__main__":
    main()
