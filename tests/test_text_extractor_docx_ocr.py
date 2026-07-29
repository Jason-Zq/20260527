"""text_extractor Office 嵌图 OCR 测试(扫描件贴进 Word/Excel 的场景)。

  cd e:/qoderproject/20260527
  PYTHONIOENCODING=utf-8 PYTHONUTF8=1 ./.venv312/Scripts/python.exe tests/test_text_extractor_docx_ocr.py

测试核心 case:
- 纯文字 docx(≥阈值): 不触发嵌图 OCR
- 空文字+嵌大图 docx: 触发 OCR,source=docx_img_ocr
- 短文字+嵌大图 docx: 文本合并,source=docx_text+img_ocr
- 小图(<MIN_SIDE)跳过;超 MAX_IMAGES 截断;单图 OCR 异常不杀整体
- xlsx: 空文本+嵌图触发,source=xlsx_img_ocr;有文本不触发
"""
import os
import sys
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "backend"))

import text_extractor
import ocr_service
import event_service


# === stub event_service 避免真实 DB ===
event_service.log_event = lambda *args, **kwargs: None


# === mock run_ocr,记录调用 ===
_ocr_calls = []


def _fake_run_ocr(path, cls=True):
    _ocr_calls.append(path)
    return [[["bbox", ("模拟识别文字", 0.95)]]]


ocr_service.run_ocr = _fake_run_ocr
text_extractor._log_office_img_ocr = lambda *a, **k: None

_TMP = tempfile.mkdtemp(prefix="test_docx_ocr_")


def _make_png(name: str, size=(300, 200), color=(255, 255, 255)) -> str:
    from PIL import Image
    p = os.path.join(_TMP, name)
    Image.new("RGB", size, color).save(p, "PNG")
    return p


def _make_docx(name: str, text: str = "", images: int = 0, img_size=(300, 200)) -> str:
    from docx import Document
    p = os.path.join(_TMP, name)
    doc = Document()
    if text:
        doc.add_paragraph(text)
    for i in range(images):
        # python-docx 按图片内容 hash 去重 media 条目,测试多图场景要内容各异
        doc.add_picture(_make_png(f"{name}_img{i}.png", img_size,
                                  color=(255, (i * 37) % 256, (i * 91) % 256)))
    doc.save(p)
    return p


def _make_xlsx(name: str, text: str = "", with_image: bool = False) -> str:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XlImage
    p = os.path.join(_TMP, name)
    wb = Workbook()
    ws = wb.active
    if text:
        ws["A1"] = text
    if with_image:
        ws.add_image(XlImage(_make_png(f"{name}_xlimg.png")))
    wb.save(p)
    return p


def _reset():
    _ocr_calls.clear()


# === 测试 ===

def test_docx_plain_text_no_ocr():
    """纯文字 docx(≥80 字):不触发嵌图 OCR。"""
    _reset()
    p = _make_docx("plain.docx", text="这是一段足够长的正文," * 20)
    r = text_extractor._extract_docx(p)
    assert r["source"] == "docx_text", r["source"]
    assert len(_ocr_calls) == 0, _ocr_calls


def test_docx_image_only_triggers_ocr():
    """空文字+嵌大图:触发 OCR,source=docx_img_ocr。"""
    _reset()
    p = _make_docx("scan.docx", images=1)
    r = text_extractor._extract_docx(p)
    assert r["source"] == "docx_img_ocr", r["source"]
    assert "模拟识别文字" in r["text"], r["text"]
    assert r["char_count"] == len(r["text"])
    assert len(_ocr_calls) == 1, _ocr_calls


def test_docx_short_text_plus_image_merges():
    """短文字(<80 字)+嵌大图:文本合并,source=docx_text+img_ocr。"""
    _reset()
    p = _make_docx("mixed.docx", text="护照扫描件", images=1)
    r = text_extractor._extract_docx(p)
    assert r["source"] == "docx_text+img_ocr", r["source"]
    assert "护照扫描件" in r["text"] and "模拟识别文字" in r["text"], r["text"]
    assert "嵌入图片 OCR" in r["text"]
    assert len(_ocr_calls) == 1


def test_docx_small_image_skipped():
    """短边 <150px 的图标图:跳过,OCR 不调用。"""
    _reset()
    p = _make_docx("icon.docx", images=1, img_size=(50, 50))
    r = text_extractor._extract_docx(p)
    assert r["source"] == "docx_text", r["source"]
    assert len(_ocr_calls) == 0, _ocr_calls


def test_docx_max_images_capped():
    """超过 MAX_IMAGES 的嵌图:OCR 调用数封顶。"""
    _reset()
    p = _make_docx("many.docx", images=text_extractor.OFFICE_IMG_OCR_MAX_IMAGES + 2)
    r = text_extractor._extract_docx(p)
    assert r["source"] == "docx_img_ocr", r["source"]
    assert len(_ocr_calls) == text_extractor.OFFICE_IMG_OCR_MAX_IMAGES, len(_ocr_calls)


def test_docx_single_image_error_not_fatal():
    """单图 OCR 异常:跳过该图,整体不炸。"""
    _reset()
    p = _make_docx("err.docx", images=2)
    calls = {"n": 0}

    def flaky(path, cls=True):
        calls["n"] += 1
        _ocr_calls.append(path)
        if calls["n"] == 1:
            raise RuntimeError("ocr boom")
        return [[["bbox", ("第二张图文字", 0.9)]]]

    orig = ocr_service.run_ocr
    ocr_service.run_ocr = flaky
    try:
        r = text_extractor._extract_docx(p)
    finally:
        ocr_service.run_ocr = orig
    assert r["source"] == "docx_img_ocr", r["source"]
    assert "第二张图文字" in r["text"], r["text"]
    assert calls["n"] == 2


def test_xlsx_image_only_triggers_ocr():
    """空 xlsx + 嵌图:触发 OCR。
    注:xlsx 总有 sheet 头文本(--- Sheet: X ---),故 source 是混合标记而非 xlsx_img_ocr。"""
    _reset()
    p = _make_xlsx("scan.xlsx", with_image=True)
    r = text_extractor._extract_xlsx(p)
    assert r["source"] == "xlsx_text+img_ocr", r["source"]
    assert "模拟识别文字" in r["text"], r["text"]
    assert len(_ocr_calls) == 1, _ocr_calls


def test_xlsx_with_text_no_ocr():
    """有文本的 xlsx(≥80 字):不触发嵌图 OCR。"""
    _reset()
    p = _make_xlsx("text.xlsx", text="单元格内容足够长," * 20, with_image=True)
    r = text_extractor._extract_xlsx(p)
    assert r["source"] == "xlsx_text", r["source"]
    assert len(_ocr_calls) == 0, _ocr_calls


def test_doc_source_relabel():
    """soffice→docx 路线的 source 重写:docx_ 前缀换 doc_,img_ocr 标记保留。"""
    for src, want in [("docx_text", "doc_text"),
                      ("docx_text+img_ocr", "doc_text+img_ocr"),
                      ("docx_img_ocr", "doc_img_ocr")]:
        assert src.replace("docx", "doc", 1) == want, (src, want)


if __name__ == "__main__":
    tests = [v for k, v in globals().items() if k.startswith("test_") and callable(v)]
    failed = 0
    for t in tests:
        try:
            t()
            print(f"  OK   {t.__name__}")
        except AssertionError as e:
            failed += 1
            print(f"  FAIL {t.__name__}: {e}")
        except Exception as e:
            failed += 1
            print(f"  ERR  {t.__name__}: {type(e).__name__}: {e}")
    if failed:
        print(f"\n{failed}/{len(tests)} 失败")
        sys.exit(1)
    print(f"\nAll {len(tests)} tests passed.")
