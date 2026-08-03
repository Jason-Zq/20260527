# -*- coding: utf-8 -*-
"""一次性脚本：生成《云盘留底检测系统-CEO汇报.docx》。

截图/数据占位用 红色加粗 + 黄色高亮 标记，方便汇报人替换。
运行：PYTHONIOENCODING=utf-8 PYTHONUTF8=1 ./.venv312/Scripts/python.exe gen_ceo_report.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn

RED = RGBColor(0xC0, 0x00, 0x00)
DARK = RGBColor(0x1F, 0x3B, 0x63)
GRAY = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x00, 0x00, 0x00)

OUT = "云盘留底检测系统-CEO汇报.docx"


def _fmt(run, size=11, bold=False, color=None, cn="等线", en="Calibri", highlight=None):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if highlight is not None:
        run.font.highlight_color = highlight
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)


def title(doc, text, sub=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _fmt(p.add_run(text), size=20, bold=True, cn="等线")
    if sub:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        _fmt(p2.add_run(sub), size=12, color=GRAY)


def meta_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    _fmt(p.add_run(text), size=11, color=GRAY)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    _fmt(p.add_run(text), size=14, bold=True, color=DARK, cn="等线")


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    _fmt(p.add_run(text), size=12, bold=True, color=BLACK, cn="等线")


def body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    _fmt(p.add_run(text), size=11)


def bullet(doc, text, bold_head=None):
    """bold_head 传「小标题」时，小标题加粗、冒号后正文常规。"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_head:
        _fmt(p.add_run(bold_head + "："), size=11, bold=True)
    _fmt(p.add_run(text), size=11)


def marker(doc, label, hint):
    """截图/数据占位标记：红色加粗 + 黄色高亮，后接灰色小字说明。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _fmt(p.add_run(label), size=11, bold=True, color=RED, highlight=WD_COLOR_INDEX.YELLOW)
    if hint:
        _fmt(p.add_run("　" + hint), size=9, color=GRAY)


def set_cell(cell, text, bold=False, size=9, color=None):
    cell.paragraphs[0].text = ""
    _fmt(cell.paragraphs[0].add_run(text), size=size, bold=bold, color=color)


def make_table(doc, headers, rows, widths_cm):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell(tbl.rows[0].cells[j], h, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell(tbl.rows[i].cells[j], val)
    for i, r in enumerate(tbl.rows):
        for j, w in enumerate(widths_cm):
            r.cells[j].width = Cm(w)
    # 表后留点空
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.2)
        sec.left_margin = sec.right_margin = Cm(2.5)

    title(doc, "云盘留底检测系统工作汇报", "—— 系统能力 · 阶段性优化成果 · 运行成效与规划 ——")
    meta_line(doc, "汇报人：【姓名】　　2026 年 7 月")

    # 一、背景与业务痛点
    h1(doc, "一、背景与业务痛点")
    body(doc, "按公司与 PSD 业务标准，客户服务的每个关键进展（签约、递案、获批、开户、完成投资等）都必须“留底”——"
              "在云盘中留存能证明该项服务确已启动或发生的材料。材料形式繁杂：合同、受理回执、聊天记录、转款凭证、"
              "公证文件、外文材料等，此前靠人工逐份打开核对。")
    bullet(doc, "文件量大、格式杂，人工逐份核对耗时长，成本高；", bold_head="量大")
    bullet(doc, "判定口径依赖个人经验，松紧不一，结论难以复核；", bold_head="标准不一")
    bullet(doc, "业务量增长后容易漏检，合规风险后置暴露。", bold_head="易漏检")
    body(doc, "云盘留底检测系统的目标：把留底检测从“人工核对”升级为“AI 自动判定”——标准统一、全程留痕、结果可复核。")

    # 二、系统是什么
    h1(doc, "二、系统是什么")
    body(doc, "一句话：业务方把「客户 + 项目 + 进展 + 云盘文件链接」批量提交给系统，系统自动完成下载、文字识别（OCR）、"
              "AI 按公司留底标准逐份判定，并给出整批“符合 / 部分符合 / 不符合”的总体结论与理由。")
    body(doc, "处理流程：批量提交（秒级受理）→ 后台排队异步处理 → OCR 文字识别 → AI 逐份判定 → 识别文本脱敏存档 → "
              "生成批次总报告 → 业务方按批次号查询结果。同一进展下同一文件再次提交时自动复用历史判定，不重复识别、不重复计费。")
    marker(doc, "【截图 1：业务审核提交页】", "系统「业务审核」tab，框选客户/项目/进展 + 文件链接提交表单区域")
    marker(doc, "【截图 2：批次结果页】", "含单文件判定列表（结论/符合度/理由）与整批总体报告")

    # 三、核心功能
    h1(doc, "三、核心功能")
    h2(doc, "1. 业务审核（主线，与业务系统对接）")
    bullet(doc, "云盘/OSS 文件链接批量提交，秒级返回批次号，检测在后台异步完成；文件链接过期自动刷新；")
    bullet(doc, "结果含每份文件的判定结论、符合度评分（0–100）、判定理由、关键点，以及整批总体结论。")
    h2(doc, "2. 快速检测（自助工具）")
    bullet(doc, "运营/文案本地上传或粘贴链接、自定义判定要求，即传即测，用于零散核查与规则验证。")
    h2(doc, "3. 检测批次管理后台")
    bullet(doc, "批次检索：按状态、客户、项目、进展、日期、总体结论筛选，支持“仅看有失败文件”；", bold_head="检索")
    bullet(doc, "批次详情：客户/办理人/项目/进展/判定标准/总体结论全量展示，逐文件可查判定理由与识别文本；", bold_head="详情")
    bullet(doc, "单批重审；判定口径调整后可“按新规则批量重判”（不重新识别原件，成本低）；必要时批量重审。",
           bold_head="纠错机制")
    marker(doc, "【截图 3：管理后台-批次列表】", "检测批次管理页（/archive-admin），含筛选区与批次列表")
    marker(doc, "【截图 4：批次详情弹窗】", "管理后台点任一批次「详情」，展示批次全量信息 + 文件表")
    marker(doc, "【截图 5：单文件详情】", "批次详情中点某文件，展示判定理由、关键点与脱敏识别文本")

    # 四、智能判定标准
    h1(doc, "四、智能判定标准：PSD 业务标准固化进 AI")
    bullet(doc, "留底的本质是“证明该进展对应的服务确已启动或发生”，不要求必须是官方回执类关键件；", bold_head="核心口径")
    bullet(doc, "官方文件（批复函/受理回执/递交确认）或软证据组合（服务合同 + 聊天记录/邮件/确认截图）任一即可；",
           bold_head="证据认定")
    bullet(doc, "符合 / 部分符合 / 不符合三级结论 + 符合度评分 + 判定理由；只有辅助附件判“部分符合”，完全无服务痕迹判“不符合”；",
           bold_head="三级判定")
    bullet(doc, "转款凭证以金额币种为准、不强制体现客户姓名；银行开户类核心材料属客户隐私，识别到账户/开户/转款凭证任一即算符合；"
                "进展阶段错配最多判“部分符合”不硬性否决；单份文件加密或识别失败只影响该文件、不拖垮整批；",
           bold_head="业务规则示例")
    bullet(doc, "扫描件无有效文字时自动标记“无文字”，不参与总体判定，避免误判。", bold_head="异常兜底")

    # 五、阶段性优化成果
    h1(doc, "五、阶段性优化成果：提速 · 反馈闭环 · 规则迭代")
    h2(doc, "5.1 基础设施升级，效率提升 10 倍")
    body(doc, "切换到更高配置服务器后，单批次检测耗时由 160 分钟降至 12 分钟，效率提升至少 10 倍，"
              "让“改一次规则、当天跑一轮回归验证”成为常态，大幅缩短反馈闭环。")
    h2(doc, "5.2 双轨制质量监控")
    bullet(doc, "每日主动抽查检测结果，识别明显偏差案例；", bold_head="内部自检")
    bullet(doc, "对接 PSD 侧判定结论，聚焦“判定符合”背后的规则缺口。", bold_head="PSD 反馈")
    h2(doc, "5.3 已完成的规则改进（典型案例）")
    make_table(
        doc,
        ["#", "反馈场景", "PSD 判定符合的原因", "改进动作"],
        [
            ["1", "服务启动判定", "客户合同已写明具体服务内容，完成任意一项即为服务启动；聊天记录为客户确认凭证（转卖房产 = 服务启动）", "已按此规则更新判定逻辑"],
            ["2", "授权委托书识别（希腊团聚项目）", "香港公证 + 海牙认证 POA，第三页委托人信息中含项目办理人（关晓晓）", "新增规则：客户名称 + 办理人中英文名组合匹配"],
            ["3", "泰国精英签青铜卡转款凭证", "泰币 650000 已符合投资金额标准，凭证无需强制体现客户姓名；且金额已脱敏", "更新规则：不再强制客户姓名匹配，后续不比对金额数字"],
            ["4", "美国银行个户项目", "账户类材料属客户隐私，核心材料可能未直接提供给文案（开户成功后银行直接邮件联系客户）", "更新规则：成功留底中识别到客户账户信息/开通/转款凭证即为符合"],
            ["5", "新加坡银行个户项目", "同上，递交后银行直接与客户联系，文案侧无直接材料", "更新规则：递交留底中识别到开户信息表/核心个人文件等即为符合"],
            ["6", "家属关系识别（叶春燕/周小宝小白条案例）", "入境激活小白条文件名为「先生小白条」，说明周小宝为叶春燕丈夫；递交留底已上传夫妻证明文件", "新增规则：结合当前批次多进展文件整体分析 + 文件内容关键词识别"],
        ],
        [0.8, 3.4, 6.6, 5.2],
    )
    h2(doc, "5.4 已识别的边界问题与对策")
    make_table(
        doc,
        ["场景", "现状", "处理策略"],
        [
            ["加密 PDF（如香港资本投资者入境计划-完成投资批次）", "曾导致整批误判为“不符合”", "已改为：识别不了的加密文件排除出整体判断，由其他文件综合兜底"],
            ["土耳其语等小语种文件（无英文）", "AI 目前原生支持中英文，未能识别", "待适配：规划多语言识别扩展"],
        ],
        [4.6, 5.0, 6.4],
    )
    h2(doc, "5.5 调优路径")
    body(doc, "整体遵循“算力先行 → 双轨反馈 → 规则细化 → 边界补齐”：单文件判定升级为多进展、多文件联合判定；"
              "名称精确匹配升级为客户名 + 办理人中英文组合匹配；金额/姓名硬校验按项目场景放宽；"
              "异常样本降级处理，避免污染整体结论。")
    body(doc, "一句话总结：10 倍提速让每一条 PSD 反馈都能快速转化为规则更新，并通过每日自检验证效果，"
              "形成“提速 → 反馈 → 迭代 → 验证”的良性闭环。")

    # 六、技术保障与安全合规
    h1(doc, "六、技术保障与安全合规")
    bullet(doc, "队列化异步架构：提交即回、后台排队处理，集中提交不宕机；任务状态全部落库，服务重启不丢任务；",
           bold_head="稳定")
    bullet(doc, "历史结果复用 + 按新规则重判不重新识别原件，AI 调用量与成本可控；", bold_head="降本")
    bullet(doc, "事件流、接口请求记录、外部接口记录、AI 调用记录四类台账，每次判定可追溯、可复盘；", bold_head="可观测")
    bullet(doc, "识别文本脱敏后入库；文件走云盘临时签名链接、处理完即删，服务器不长期留存原件；业务接口由网络层隔离。",
           bold_head="安全")
    marker(doc, "【截图 6：事件流页】", "侧边栏「事件流」（/events），展示批次/文件/worker 事件台账")
    marker(doc, "【截图 7：AI 调用记录页】", "侧边栏「AI 调用记录」（/ai-api-calls），展示每次 AI 判定留痕")
    marker(doc, "【截图 8（可选）：在线接口文档】", "服务地址 /docs（Swagger），体现系统对外开放对接能力")

    # 七、运行成效
    h1(doc, "七、运行成效")
    body(doc, "以下数据从生产环境统计后填入：", indent=True)
    bullet(doc, "累计检测：【数据：批次 __ 批 / 文件 __ 份】；")
    bullet(doc, "总体结论分布：符合【数据：__%】、部分符合【数据：__%】、不符合【数据：__%】；")
    bullet(doc, "历史结果复用率【数据：__%】，单批次平均耗时约【数据：__】分钟；")
    bullet(doc, "内部自检 + PSD 反馈累计推动规则改进【数据：__】项。")
    marker(doc, "【数据占位说明】", "以上【数据：…】处请用生产库统计替换（批次数/文件数/结论分布/复用率/平均耗时）")

    # 八、后续规划
    h1(doc, "八、后续规划")
    bullet(doc, "同一客户同一进展多版本留痕，随材料补充自动更新总体结论；", bold_head="进展包滚动总报告")
    bullet(doc, "适配土耳其语等小语种材料，拓展识别能力边界；", bold_head="多语言扩展")
    bullet(doc, "检测完成的识别文本自动沉淀为客户结构化档案（只补空字段、不覆盖人工数据）；", bold_head="客户档案联动")
    bullet(doc, "检测结果自动回写业务系统，减少人工查询环节。", bold_head="自动回写")

    # 附录：截图清单
    doc.add_page_break()
    h1(doc, "附录：截图清单")
    make_table(
        doc,
        ["编号", "截图内容", "截取位置"],
        [
            ["截图 1", "业务审核提交页", "系统首页「业务审核」tab（提交表单区域）"],
            ["截图 2", "批次结果页（单文件判定 + 总体报告）", "提交后的批次结果轮询页"],
            ["截图 3", "管理后台-批次列表（含筛选区）", "侧边栏「检测批次管理」（/archive-admin）"],
            ["截图 4", "批次详情弹窗", "管理后台点任一批次「详情」"],
            ["截图 5", "单文件详情（判定理由 + 关键点）", "批次详情弹窗中点某文件「详情」"],
            ["截图 6", "事件流页", "侧边栏「事件流」（/events）"],
            ["截图 7", "AI 调用记录页", "侧边栏「AI 调用记录」（/ai-api-calls）"],
            ["截图 8（可选）", "在线接口文档", "服务地址 /docs（Swagger UI）"],
        ],
        [2.2, 6.4, 7.4],
    )

    doc.save(OUT)
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    main()
