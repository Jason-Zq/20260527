"""
anchor.py 单测 —— 19 个边界用例 + POA 信息表.docx 集成抽测 + 性能基准。

运行：
    cd e:/qoderproject/20260527
    PYTHONIOENCODING=utf-8 ./.venv312/Scripts/python.exe tests/test_anchor.py

不依赖 pytest，直接 python 执行。每个 test_* 函数失败抛异常，runner 汇总通过率。
"""

from __future__ import annotations

import io
import os
import sys
import time
import traceback
from copy import deepcopy

# 确保 backend/ 在 sys.path，导入 anchor 模块
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from pydantic import ValidationError

from anchor import (
    AnchorCell,
    AnchorNotFound,
    AnchorParagraph,
    AnchorRun,
    ContainerBody,
    ContainerCell,
    ContainerFooter,
    ContainerHeader,
    apply_value,
    format_value,
    locate_anchor,
)


# ====================== docx 构造助手 ======================


def make_docx_with_runs(runs_spec: list[tuple[str, dict]]):
    """构造一个含单段落的 docx，段落里有多个 run，每个 run 带指定格式。

    runs_spec: [("text", {"bold": True, "underline": True, ...}), ...]
    """
    doc = Document()
    para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    for text, fmt in runs_spec:
        run = para.add_run(text)
        for key, value in fmt.items():
            setattr(run, key, value)
    return doc


def make_docx_with_table(rows: int = 2, cols: int = 2, fill: dict[tuple[int, int], str] | None = None):
    """构造一个含单表格的 docx。fill 用于预填某些 cell 文本。"""
    fill = fill or {}
    doc = Document()
    table = doc.add_table(rows=rows, cols=cols)
    for ri in range(rows):
        for ci in range(cols):
            text = fill.get((ri, ci), "")
            table.rows[ri].cells[ci].text = text
    return doc


def make_docx_with_merged_cells():
    """构造含合并 cell 的 docx：2x2 表格，第一行横向合并。"""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # 第一行 c0 + c1 合并
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    table.rows[1].cells[0].text = "left"
    table.rows[1].cells[1].text = "right"
    return doc


def xml_bytes_of(elem) -> bytes:
    """把 lxml 元素序列化成字节用于对比。"""
    return etree.tostring(elem, encoding="utf-8")


def docx_xml_dump(doc) -> bytes:
    """把整个 doc.element 序列化为字节。"""
    return etree.tostring(doc.element.body, encoding="utf-8")


# ====================== 19 个边界用例 ======================


def test_01_empty_cell_write_value():
    """用例 1：空 cell 写值。期望：cell 出现一个 run，cell 边框/底纹/列宽不变。"""
    doc = make_docx_with_table(rows=2, cols=2)
    cell = doc.tables[0].rows[0].cells[0]
    # 记录 tcPr 原始 XML（边框/底纹/对齐都在这里面）
    tc_pr_before = cell._tc.find(qn("w:tcPr"))
    tc_pr_before_xml = xml_bytes_of(tc_pr_before) if tc_pr_before is not None else None

    apply_value(doc, AnchorCell(kind="cell", t=0, r=0, c=0), "Hello")

    # cell.text 应该有值
    assert cell.text == "Hello", f"cell.text 期望 'Hello' 实际 {cell.text!r}"
    # tcPr 应完全不变
    tc_pr_after = cell._tc.find(qn("w:tcPr"))
    tc_pr_after_xml = xml_bytes_of(tc_pr_after) if tc_pr_after is not None else None
    assert tc_pr_before_xml == tc_pr_after_xml, "tcPr XML 被修改了"
    # 其他 cell 不应被动
    assert doc.tables[0].rows[0].cells[1].text == "", "兄弟 cell 被改了"
    assert doc.tables[0].rows[1].cells[0].text == "", "下一行 cell 被改了"


def test_02_merged_cell_write_top_left():
    """用例 2：合并 cell 写左上起点。期望：合并范围保留。"""
    doc = make_docx_with_merged_cells()
    apply_value(doc, AnchorCell(kind="cell", t=0, r=0, c=0), "Merged")
    # 第一行左上应有 'Merged'
    assert doc.tables[0].rows[0].cells[0].text == "Merged"
    # 合并 cell 在 python-docx 中重复返回，c=1 应当也是同样的 'Merged'
    # （因为 _is_merged_continuation 检测后会拒绝写 c=1）
    assert doc.tables[0].rows[0].cells[1].text == "Merged", \
        "合并 cell 的镜像应同步显示同值"
    # 第二行不动
    assert doc.tables[0].rows[1].cells[0].text == "left"
    assert doc.tables[0].rows[1].cells[1].text == "right"


def test_03_merged_cell_continuation_rejected():
    """用例 3：cell anchor 指向合并 cell 的被合并位，期望 AnchorNotFound。"""
    doc = make_docx_with_merged_cells()
    try:
        apply_value(doc, AnchorCell(kind="cell", t=0, r=0, c=1), "X")
    except AnchorNotFound as e:
        assert "合并" in str(e), f"错误信息应提到合并: {e}"
        return
    raise AssertionError("AnchorNotFound 未抛出")


def test_04_paragraph_full_replace():
    """用例 4：paragraph 整段替换。期望：段内只剩 1 run，沿用首 run rPr，pPr 不动。"""
    doc = make_docx_with_runs([
        ("姓名：", {"bold": True}),
        ("__",    {"underline": True}),
        (" 先生", {}),
    ])
    # 记录 pPr 原始 XML
    para = doc.paragraphs[0]
    p_pr_before = para._p.find(qn("w:pPr"))
    p_pr_before_xml = xml_bytes_of(p_pr_before) if p_pr_before is not None else None

    apply_value(doc, AnchorParagraph(kind="paragraph", container=ContainerBody(kind="body"), p=0), "王二")

    runs = para.runs
    assert len(runs) == 1, f"段内 run 数期望 1 实际 {len(runs)}"
    assert runs[0].text == "王二"
    assert runs[0].bold is True, "应沿用首 run 的 bold"
    # pPr 不变
    p_pr_after = para._p.find(qn("w:pPr"))
    p_pr_after_xml = xml_bytes_of(p_pr_after) if p_pr_after is not None else None
    assert p_pr_before_xml == p_pr_after_xml, "pPr XML 被修改了"


def test_05_run_full_replace_preserves_siblings():
    """用例 5：run 整 run 替换。期望：该 run.text 变；其他 run 不动；rPr 完全不动。"""
    doc = make_docx_with_runs([
        ("姓名：", {"bold": True}),
        ("__",    {"underline": True}),
        (" 先生", {}),
    ])
    para = doc.paragraphs[0]
    # 记录 run[0] / run[2] 的 XML
    run0_xml_before = xml_bytes_of(para.runs[0]._r)
    run2_xml_before = xml_bytes_of(para.runs[2]._r)
    # 记录 run[1] 的 rPr
    run1_rpr_before = xml_bytes_of(para.runs[1]._r.find(qn("w:rPr")))

    apply_value(
        doc,
        AnchorRun(kind="run", container=ContainerBody(kind="body"), p=0, run_index=1),
        "王二",
    )

    # 兄弟 run 完全不变
    assert xml_bytes_of(para.runs[0]._r) == run0_xml_before, "run[0] 被动"
    assert xml_bytes_of(para.runs[2]._r) == run2_xml_before, "run[2] 被动"
    # 目标 run text 变了
    assert para.runs[1].text == "王二"
    # 目标 run rPr 不变
    run1_rpr_after = xml_bytes_of(para.runs[1]._r.find(qn("w:rPr")))
    assert run1_rpr_after == run1_rpr_before, "目标 run rPr 被改"
    # underline 仍在
    assert para.runs[1].underline is True


def test_06_run_char_slice_replace():
    """用例 6：run 内字符级替换。原 text='姓名：XX 先生'，offset=3,length=2 → '姓名：王二 先生'。"""
    doc = make_docx_with_runs([("姓名：XX 先生", {"bold": True})])
    para = doc.paragraphs[0]
    rpr_before = xml_bytes_of(para.runs[0]._r.find(qn("w:rPr")))

    apply_value(
        doc,
        AnchorRun(
            kind="run",
            container=ContainerBody(kind="body"),
            p=0,
            run_index=0,
            char_offset=3,
            char_length=2,
        ),
        "王二",
    )

    assert para.runs[0].text == "姓名：王二 先生", f"实际: {para.runs[0].text!r}"
    # rPr 完全不动
    rpr_after = xml_bytes_of(para.runs[0]._r.find(qn("w:rPr")))
    assert rpr_after == rpr_before
    assert para.runs[0].bold is True


def test_07_run_char_slice_out_of_range():
    """用例 7：run 内 offset 越界，期望 AnchorNotFound。"""
    doc = make_docx_with_runs([("short", {})])
    try:
        apply_value(
            doc,
            AnchorRun(
                kind="run",
                container=ContainerBody(kind="body"),
                p=0,
                run_index=0,
                char_offset=3,
                char_length=10,  # 3+10 > 5
            ),
            "X",
        )
    except AnchorNotFound as e:
        assert "too short" in str(e) or "short" in str(e), f"错误信息异常: {e}"
        return
    raise AssertionError("AnchorNotFound 未抛出")


def test_08_paragraph_replace_clears_hyperlink():
    """用例 8：段内含 hyperlink 的 run，paragraph 整段替换会清掉 hyperlink。
    （文档化为预期：用户要保 hyperlink 应当用 run anchor 精确定位）"""
    doc = make_docx_with_runs([("Normal text", {})])
    para = doc.paragraphs[0]
    # 程序化插入一个 hyperlink 节点
    hl = etree.SubElement(para._p, qn("w:hyperlink"))
    hl_run = etree.SubElement(hl, qn("w:r"))
    hl_text = etree.SubElement(hl_run, qn("w:t"))
    hl_text.text = "Link"
    # 此时段落里有 1 run + 1 hyperlink

    apply_value(doc, AnchorParagraph(kind="paragraph", container=ContainerBody(kind="body"), p=0), "Replaced")

    # hyperlink 被清掉是预期
    assert para._p.find(qn("w:hyperlink")) is None, "hyperlink 应被清掉"
    # 只剩新加的 1 run
    assert len(para.runs) == 1
    assert para.runs[0].text == "Replaced"


def test_09_header_paragraph_replace():
    """用例 9：替换 header 首段，body 不动。"""
    doc = Document()
    # 加段落到 body
    doc.add_paragraph("body para")
    # 加内容到 header
    section = doc.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.add_run("HEADER")

    body_xml_before = docx_xml_dump(doc)

    apply_value(
        doc,
        AnchorParagraph(kind="paragraph", container=ContainerHeader(kind="header", s=0), p=0),
        "NEW HEADER",
    )

    # body 一字不变
    assert docx_xml_dump(doc) == body_xml_before, "body 被修改"
    # header 变了
    assert section.header.paragraphs[0].text == "NEW HEADER"


def test_10_multi_section_footer_replace():
    """用例 10：多 section 的 footer，仅第 2 section 变。"""
    doc = Document()
    # 第 1 个 section 的 footer
    doc.sections[0].footer.paragraphs[0].add_run("FOOTER 1")
    # 加第 2 个 section
    from docx.enum.section import WD_SECTION
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.footer.is_linked_to_previous = False  # 解除继承
    new_section.footer.paragraphs[0].add_run("FOOTER 2")

    # 替换第 2 section 的 footer
    apply_value(
        doc,
        AnchorParagraph(kind="paragraph", container=ContainerFooter(kind="footer", s=1), p=0),
        "NEW FOOTER 2",
    )

    # 第 1 section 的 footer 不变
    assert doc.sections[0].footer.paragraphs[0].text == "FOOTER 1", \
        f"section 0 footer 应为 'FOOTER 1' 实际 {doc.sections[0].footer.paragraphs[0].text!r}"
    # 第 2 section 的 footer 变了
    assert doc.sections[1].footer.paragraphs[0].text == "NEW FOOTER 2"


def test_11_paragraph_in_cell_container():
    """用例 11：cell 容器内的段落替换。cell 边框不动。"""
    doc = make_docx_with_table(rows=2, cols=2, fill={(0, 0): "原始内容"})
    cell = doc.tables[0].rows[0].cells[0]
    tc_pr_before = xml_bytes_of(cell._tc.find(qn("w:tcPr"))) if cell._tc.find(qn("w:tcPr")) is not None else None

    apply_value(
        doc,
        AnchorParagraph(
            kind="paragraph",
            container=ContainerCell(kind="cell", t=0, r=0, c=0),
            p=0,
        ),
        "新内容",
    )

    assert cell.text == "新内容"
    tc_pr_after = xml_bytes_of(cell._tc.find(qn("w:tcPr"))) if cell._tc.find(qn("w:tcPr")) is not None else None
    assert tc_pr_before == tc_pr_after, "tcPr 被改"


def test_12_date_format_iso():
    """用例 12：date fmt YYYY-MM-DD。"""
    from datetime import datetime
    s = format_value(datetime(2026, 6, 11), "date:YYYY-MM-DD")
    assert s == "2026-06-11", f"实际: {s!r}"


def test_13_date_format_chinese():
    """用例 13：date fmt 中文。"""
    from datetime import date
    s = format_value(date(2026, 6, 11), "date:YYYY年MM月DD日")
    assert s == "2026年06月11日", f"实际: {s!r}"


def test_14_currency_format():
    """用例 14：currency fmt。"""
    s = format_value(12345.67, "currency:CNY")
    assert s == "12,345.67 CNY", f"实际: {s!r}"


def test_15_none_value():
    """用例 15：value=None 写入空串。"""
    doc = make_docx_with_runs([("X", {})])
    apply_value(
        doc,
        AnchorRun(kind="run", container=ContainerBody(kind="body"), p=0, run_index=0),
        None,
    )
    assert doc.paragraphs[0].runs[0].text == ""


def test_16_special_chars_in_value():
    """用例 16：value 含 < > & 字面保留，xml escape 自动处理。"""
    doc = make_docx_with_runs([("X", {})])
    apply_value(
        doc,
        AnchorRun(kind="run", container=ContainerBody(kind="body"), p=0, run_index=0),
        "A < B & C > D",
    )
    assert doc.paragraphs[0].runs[0].text == "A < B & C > D"
    # 序列化 + 反序列化应一致（验证 lxml 正确 escape）
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    assert doc2.paragraphs[0].runs[0].text == "A < B & C > D"


def test_17_validator_rejects_half_char_args():
    """用例 17：char_offset 给了但 char_length 没给，Pydantic 拒绝。"""
    try:
        AnchorRun(
            kind="run",
            container=ContainerBody(kind="body"),
            p=0,
            run_index=0,
            char_offset=3,
        )
    except ValidationError as e:
        assert "char_offset" in str(e) or "char_length" in str(e), f"错误信息: {e}"
        return
    raise AssertionError("ValidationError 未抛出")


def test_18_header_section_out_of_range():
    """用例 18：header.s 越界，AnchorNotFound。"""
    doc = Document()
    try:
        apply_value(
            doc,
            AnchorParagraph(kind="paragraph", container=ContainerHeader(kind="header", s=99), p=0),
            "X",
        )
    except AnchorNotFound:
        return
    raise AssertionError("AnchorNotFound 未抛出")


def test_19_two_runs_in_same_paragraph():
    """用例 19：连续 apply_value 到同段不同 run，中间 run 不动。"""
    doc = make_docx_with_runs([
        ("A", {"bold": True}),
        ("B", {"italic": True}),
        ("C", {}),
    ])
    para = doc.paragraphs[0]
    mid_xml_before = xml_bytes_of(para.runs[1]._r)

    apply_value(
        doc,
        AnchorRun(kind="run", container=ContainerBody(kind="body"), p=0, run_index=0),
        "AA",
    )
    apply_value(
        doc,
        AnchorRun(kind="run", container=ContainerBody(kind="body"), p=0, run_index=2),
        "CC",
    )

    assert para.runs[0].text == "AA" and para.runs[0].bold is True
    assert para.runs[1].text == "B"
    assert xml_bytes_of(para.runs[1]._r) == mid_xml_before, "中间 run 被动"
    assert para.runs[2].text == "CC"


# ====================== POA 信息表.docx 集成抽测 ======================


def test_integration_poa_form():
    """用 POA 信息表.docx 真实业务模板做集成测试。

    检查：
      - 4 个表格的"姓名"行 c=1 是空 cell
      - 写入 4 个不同值后，未触碰 cell 的 XML 字节级不变
      - 序列化后能再次打开
    """
    src = os.path.join(os.path.dirname(__file__), "..", "POA 信息表.docx")
    if not os.path.exists(src):
        raise AssertionError(f"POA 信息表.docx 不存在: {src}")

    doc = Document(src)
    assert len(doc.tables) == 4, f"POA 应有 4 个表格，实际 {len(doc.tables)}"

    # 记录所有 cell 的 XML 用于对比"未触碰 cell 不变"
    before_snapshots = []
    for ti, table in enumerate(doc.tables):
        snap = []
        for ri, row in enumerate(table.rows):
            row_snap = []
            seen = set()
            for ci, cell in enumerate(row.cells):
                if id(cell._tc) in seen:
                    row_snap.append(None)  # 合并 cell 占位
                else:
                    seen.add(id(cell._tc))
                    row_snap.append(xml_bytes_of(cell._tc))
            snap.append(row_snap)
        before_snapshots.append(snap)

    # 写 4 个表的"Name/姓名"行 c=1
    names = ["主申请人", "随行家属甲", "随行家属乙", "随行家属丙"]
    for ti, name in enumerate(names):
        # POA 模板第 1 行（r=1）是 Name/姓名
        apply_value(doc, AnchorCell(kind="cell", t=ti, r=1, c=1), name)

    # 验证写入成功
    for ti, name in enumerate(names):
        actual = doc.tables[ti].rows[1].cells[1].text
        assert actual == name, f"表{ti} r=1 c=1 期望 {name!r} 实际 {actual!r}"

    # 验证未触碰 cell XML 不变
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen = set()
            for ci, cell in enumerate(row.cells):
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                if (ri, ci) == (1, 1):
                    continue  # 这是被改的 cell
                before = before_snapshots[ti][ri][ci]
                if before is None:
                    continue
                after = xml_bytes_of(cell._tc)
                assert before == after, (
                    f"未触碰 cell (t={ti},r={ri},c={ci}) 被修改:\n"
                    f"before: {before[:200]}\n"
                    f"after:  {after[:200]}"
                )

    # 验证可序列化重读
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    doc2 = Document(buf)
    for ti, name in enumerate(names):
        assert doc2.tables[ti].rows[1].cells[1].text == name


# ====================== 性能基准 ======================


def test_perf_100x_apply_under_500ms():
    """性能基准：100 次 apply_value 在 POA 信息表（~45KB）上 < 500ms。"""
    src = os.path.join(os.path.dirname(__file__), "..", "POA 信息表.docx")
    if not os.path.exists(src):
        raise AssertionError(f"POA 信息表.docx 不存在: {src}")

    doc = Document(src)
    # 每次都对 cell (t=0, r=1, c=1) 写新值
    anchor = AnchorCell(kind="cell", t=0, r=1, c=1)

    t0 = time.perf_counter()
    for i in range(100):
        # 重新打开避免 cell 内容累积
        if i % 10 == 0:
            doc = Document(src)
        apply_value(doc, anchor, f"value_{i}")
    elapsed = time.perf_counter() - t0

    print(f"  perf: 100 次 apply_value 耗时 {elapsed * 1000:.1f}ms")
    assert elapsed < 0.5, f"超时: {elapsed:.3f}s 应 < 0.5s"


# ====================== Runner ======================


TESTS = [
    ("01_empty_cell_write_value",          test_01_empty_cell_write_value),
    ("02_merged_cell_write_top_left",      test_02_merged_cell_write_top_left),
    ("03_merged_cell_continuation_reject", test_03_merged_cell_continuation_rejected),
    ("04_paragraph_full_replace",          test_04_paragraph_full_replace),
    ("05_run_full_replace",                test_05_run_full_replace_preserves_siblings),
    ("06_run_char_slice_replace",          test_06_run_char_slice_replace),
    ("07_run_char_slice_out_of_range",     test_07_run_char_slice_out_of_range),
    ("08_paragraph_replace_clears_hl",     test_08_paragraph_replace_clears_hyperlink),
    ("09_header_paragraph_replace",        test_09_header_paragraph_replace),
    ("10_multi_section_footer_replace",    test_10_multi_section_footer_replace),
    ("11_paragraph_in_cell_container",     test_11_paragraph_in_cell_container),
    ("12_date_format_iso",                 test_12_date_format_iso),
    ("13_date_format_chinese",             test_13_date_format_chinese),
    ("14_currency_format",                 test_14_currency_format),
    ("15_none_value",                      test_15_none_value),
    ("16_special_chars_in_value",          test_16_special_chars_in_value),
    ("17_validator_rejects_half_args",     test_17_validator_rejects_half_char_args),
    ("18_header_section_out_of_range",     test_18_header_section_out_of_range),
    ("19_two_runs_in_same_paragraph",      test_19_two_runs_in_same_paragraph),
    ("integration_poa_form",               test_integration_poa_form),
    ("perf_100x_apply_under_500ms",        test_perf_100x_apply_under_500ms),
]


def main():
    passed = 0
    failed = []
    for name, fn in TESTS:
        try:
            fn()
            print(f"  PASS  {name}")
            passed += 1
        except Exception as e:
            print(f"  FAIL  {name}: {type(e).__name__}: {e}")
            traceback.print_exc()
            failed.append(name)
    total = len(TESTS)
    print(f"\n{passed}/{total} 通过")
    if failed:
        print("失败:", ", ".join(failed))
        sys.exit(1)
    sys.exit(0)


if __name__ == "__main__":
    main()
